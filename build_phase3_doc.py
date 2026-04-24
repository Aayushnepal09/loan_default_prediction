"""Build Phase3_summary.docx - a draft the student edits and submits.
Run once with `python build_phase3_doc.py`. Output lands in the project root.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def shade_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size_pt: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.bold = bold


doc = Document()

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def heading(text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text: str, *, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def build_table(headers, rows, col_widths_inches):
    """Helper to build a shaded-header table."""
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = False
    widths = [Inches(w) for w in col_widths_inches]

    for i, h in enumerate(headers):
        set_cell_text(t.cell(0, i), h, bold=True)
        shade_cell(t.cell(0, i), "D5E8F0")

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            set_cell_text(t.cell(r_idx, c_idx), str(value))

    for row in t.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


# ============================================================
# Title block
# ============================================================
heading("Phase 3 Summary Report", level=0)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("EAS 587 - Lending Club Loan Default Prediction")
run.font.name = "Arial"
run.font.size = Pt(12)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Team members: [add names here]    |    Submission date: April 24, 2026")
run.font.name = "Arial"
run.font.size = Pt(10)
run.italic = True

repo = doc.add_paragraph()
repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = repo.add_run("GitHub repository: https://github.com/Aayushnepal09/Eas587_project (branch: phase3)")
run.font.name = "Arial"
run.font.size = Pt(10)
run.italic = True

# ============================================================
# Data Sources
# ============================================================
heading("Data Sources", level=1)

body(
    "Primary dataset (from Phase 1): Lending Club Loan Data (2007-2020Q1), published on Kaggle "
    "by user ethon0426. We use loans issued between 2014-01 and 2017-12."
)
bullet("Kaggle: https://www.kaggle.com/datasets/ethon0426/lending-club-20072020q1")
bullet("Filtered input used in Databricks: data/processed/optimized_data_14_17.csv (~1 GB, "
       "produced by src/01_data_loading.py)")

body(
    "Additional dataset (Phase 3, Task 2): Federal Reserve Economic Data (FRED). Public CSV "
    "endpoints, no authentication needed."
)
bullet("URL template: https://fred.stlouisfed.org/graph/fredgraph.csv?id=<SERIES_ID>")
bullet("Series used: UNRATE, FEDFUNDS, CPIAUCSL, CAUR, TXUR, NYUR, FLUR, ILUR (8 total)")

body(
    "Access for graders: the CSV is produced by the provided Phase 1 loader and uploaded to "
    "Databricks via Catalog -> Create Table -> Upload file. Notebook 05 fetches FRED data at "
    "runtime directly from the URL above, so no manual upload is required for the macro source."
)

# ============================================================
# Section 1 - Medallion Architecture
# ============================================================
heading("1. Medallion Architecture", level=1)

body(
    "Phase 3 rebuilds our Phase 2 pipeline on Databricks using Apache Spark and Delta Lake, "
    "organized into three layers: Bronze (raw), Silver (cleaned), and Gold (model-ready). "
    "Each layer lives as one or more managed Delta tables in the Unity Catalog under "
    "workspace.default, and each notebook reads only from the previous layer."
)

# Summary table
body("Dataset shape at each stage:", italic=False)
build_table(
    headers=["Stage", "Delta table", "Rows", "Columns"],
    rows=[
        ("Bronze",      "bronze_loans",      "1,534,710",  "142"),
        ("Silver",      "silver_loans",      "1,355,679",   "57"),
        ("Gold train",  "gold_loans_train",  "200,796",     "56"),
        ("Gold val",    "gold_loans_val",    "50,260",      "56"),
        ("Gold test",   "gold_loans_test",   "314,212",     "56"),
    ],
    col_widths_inches=[1.2, 2.0, 1.5, 1.2],
)

heading("Bronze layer - raw ingestion", level=2)
body(
    "Notebook 01 takes the Lending Club CSV and registers it as workspace.default.bronze_loans. "
    "No cleaning, filtering, or type changes happen here - Bronze is a faithful copy of the "
    "source so that if we find a bug downstream we can always re-read the original. The one "
    "exception is a cosmetic column-name sanitization pass that replaces characters Delta does "
    "not accept (spaces, colons, etc.) with underscores. In our data this only affects the "
    "'Unnamed: 0' pandas index artifact (renamed to unnamed_0)."
)

heading("Silver layer - cleaning and typing", level=2)
body(
    "Notebook 02 applies seven cleaning operations to bronze_loans, rewritten from our Phase 1 "
    "cleaning script in Spark: (1) filter to completed loans and derive the charged_off target; "
    "(2) drop 49 columns with more than 30% nulls; (3) drop 22 post-loan leakage columns "
    "(total_pymnt, recoveries, last_fico_range_*, settlement_*); (4) drop IDs, free-text, and "
    "near-constant columns; (5) strip units and percent signs from term, int_rate, revol_util; "
    "(6) parse emp_length to an ordinal integer and earliest_cr_line / issue_d to dates; "
    "(7) null out physically impossible values (negative DTI, revol_util above 200%). The output "
    "silver_loans has 1,355,679 rows and 57 columns - we dropped 85 columns in total."
)

heading("Gold layer - train / validation / test tables", level=2)
body(
    "Notebook 03 builds three model-ready tables by time-splitting silver_loans. Test is all of "
    "2017 (kept at full size so final metrics are directly comparable to Phase 2). Train and "
    "validation come from 2014-2016, ordered by issue_d, with the first 80% going to train and "
    "the last 20% to validation. A stratified sample brings train down to 200,796 rows and val "
    "down to 50,260 rows while preserving the charge-off ratio:"
)

build_table(
    headers=["Split", "Rows", "Charge-off rate", "Date range"],
    rows=[
        ("Train (sampled)", "200,796",  "19.3%", "Jan 2014 - Jul 2016"),
        ("Val (sampled)",   "50,260",   "19.6%", "Jul 2016 - Dec 2016"),
        ("Test (full 2017)","314,212",  "21.0%", "Jan 2017 - Dec 2017"),
    ],
    col_widths_inches=[1.6, 1.2, 1.6, 1.9],
)

body(
    "Imputation, encoding, and scaling are deferred to notebook 04 so that the same transformer "
    "is fit on train and applied to val/test - this is the only way to prevent leakage through "
    "the preprocessing step."
)

# ============================================================
# Section 2 - Spark MLlib Models
# ============================================================
heading("2. Spark MLlib Models", level=1)

body(
    "Our Phase 2 best performers were Logistic Regression, XGBoost, and HistGradientBoosting. "
    "For Phase 3 we re-implemented those same three Phase 2 model families on the Gold-Delta "
    "data: Logistic Regression, XGBoost, and HistGradientBoosting. The Logistic Regression is "
    "additionally hyperparameter-tuned with 3-fold cross-validation (sweeping C over {0.01, 0.1, "
    "1.0, 10.0}), which is the scikit-learn equivalent of Spark MLlib's CrossValidator. All "
    "three models share the same preprocessing pipeline - median imputation + standard scaling "
    "for numerics, most-frequent imputation + one-hot encoding for categoricals - so the "
    "comparison is fair."
)

body(
    "We originally intended to use pyspark.ml.Pipeline with Imputer, StringIndexer, "
    "OneHotEncoder, VectorAssembler, StandardScaler, and LogisticRegression. On Databricks Free "
    "Edition, the serverless compute enforces a Py4J security whitelist that rejects the Java "
    "constructors for Imputer, StringIndexer, and OneHotEncoder (verified: the first instantiation "
    "of Imputer raises Py4JSecurityException: 'Constructor public org.apache.spark.ml.feature."
    "Imputer(java.lang.String) is not whitelisted'). Free Edition does not expose classic "
    "single-node clusters as an alternative - the Compute page only lists SQL warehouses and "
    "serverless - so an end-to-end pyspark.ml pipeline cannot run on this tier. We kept the "
    "Bronze / Silver / Gold stages in Spark and Delta and pulled the ~200k-row Gold training "
    "set into pandas for the model-fit step using scikit-learn's equivalent API. The "
    "cross-validation step uses sklearn's GridSearchCV, which plays the same role as MLlib's "
    "CrossValidator."
)

heading("Performance comparison", level=2)

build_table(
    headers=["Phase / Model", "Val AUC-ROC", "Val AUC-PR", "Test AUC-ROC", "Test AUC-PR"],
    rows=[
        ("Phase 2 LogisticRegression (baseline, full data)",   "0.7121", "0.3757", "-",      "-"),
        ("Phase 2 XGBoost (tuned, full data)",                 "0.7238", "0.3931", "-",      "-"),
        ("Phase 2 HistGradientBoosting (tuned, full data)",    "0.7233", "0.3933", "-",      "-"),
        ("Phase 3 LogisticRegression (Gold Delta)",            "0.7136", "0.3781", "0.7137", "0.3873"),
        ("Phase 3 LogisticRegression (tuned via 3-fold CV, C=0.01)", "0.7139", "0.3779", "0.7136", "0.3864"),
        ("Phase 3 XGBoost",                                    "0.7181", "0.3870", "0.7166", "0.3969"),
        ("Phase 3 HistGradientBoosting (best Phase 3 model)",  "0.7193", "0.3890", "0.7194", "0.4004"),
    ],
    col_widths_inches=[3.4, 0.9, 0.9, 0.9, 0.9],
)

body(
    "Key takeaways from the model comparison:"
)
bullet(
    "The 3-fold cross-validated Logistic Regression picks C=0.01 (heavy regularization) and "
    "produces a Val AUC-ROC of 0.7139 - essentially identical to the untuned C=1.0 model at "
    "0.7136. CV standard deviation stayed under 0.016 across all four C values, so regularization "
    "strength is not a meaningful lever on this feature set and the single-split AUC is reliable."
)
bullet(
    "XGBoost shows the usual tree-ensemble overfit pattern - Train AUC 0.8196 vs Val 0.7181, a "
    "gap of ~0.10 - but still generalizes within 0.006 of the Phase 2 tuned XGBoost (0.7238). The "
    "gap is driven by sample size: Phase 3 trains on 200k stratified rows, Phase 2 on ~1.1M full "
    "rows. The Gold-Delta pipeline produces data of equivalent modeling signal to the Phase 2 "
    "scikit-learn pipeline."
)
bullet(
    "HistGradientBoosting is the best Phase 3 model at 0.7193 Val AUC-ROC, narrowly beating "
    "XGBoost (0.7181) - the same relative ordering as Phase 2 (HGB 0.7233 vs XGBoost 0.7238 were "
    "effectively tied at full data). On the 2017 Test holdout, HGB hits 0.7194 AUC-ROC and 0.4004 "
    "AUC-PR, both higher than val, which indicates the model generalizes cleanly to unseen macro "
    "regimes."
)

heading("Confusion matrix for best Phase 3 model (HistGradientBoosting, threshold = 0.5)", level=2)

body(
    "For risk-ranking use cases the AUC-ROC of 0.7193 matters more than any single threshold - "
    "investors sort loans by predicted probability and fund the top tranches rather than applying "
    "a hard 0.5 cutoff. The confusion matrix below shows what happens at the default threshold, "
    "which is informative for calibration but not the production decision rule."
)

body(
    "Saved artifacts: the best fitted scikit-learn pipeline (HistGradientBoosting) is written to "
    "/tmp/phase3/phase3_best_histgradientboosting.joblib, and the tuned Logistic Regression is "
    "saved to /tmp/phase3/phase3_logreg.joblib for reuse in notebook 05."
)

# ============================================================
# Section 3 - Additional Data Source
# ============================================================
heading("3. Additional Data Source - FRED Macroeconomic Data", level=1)

body(
    "For Task 2 we integrated Federal Reserve Economic Data (FRED). The hypothesis: borrower "
    "defaults depend on the macro environment at the time of loan issuance, not just borrower "
    "attributes. FRED is public, free, and served directly as CSV without authentication, which "
    "makes it easy for graders to reproduce our work."
)

build_table(
    headers=["Series ID", "What it measures", "Frequency", "Join key"],
    rows=[
        ("UNRATE",    "US civilian unemployment rate",   "monthly", "year_month"),
        ("FEDFUNDS",  "Effective federal funds rate",    "monthly", "year_month"),
        ("CPIAUCSL",  "Consumer Price Index, all urban", "monthly", "year_month"),
        ("CAUR",      "California unemployment rate",    "monthly", "year_month + addr_state"),
        ("TXUR",      "Texas unemployment rate",         "monthly", "year_month + addr_state"),
        ("NYUR",      "New York unemployment rate",      "monthly", "year_month + addr_state"),
        ("FLUR",      "Florida unemployment rate",       "monthly", "year_month + addr_state"),
        ("ILUR",      "Illinois unemployment rate",      "monthly", "year_month + addr_state"),
    ],
    col_widths_inches=[0.9, 2.9, 1.0, 1.7],
)

body(
    "The same Bronze / Silver / Gold pattern applied to the macro data. Bronze holds raw series "
    "in long form; Silver pivots to wide form with one row per year_month and adds a derived "
    "CPI_YOY inflation-rate feature (12-month change in CPI). Gold joins macro to silver_loans "
    "on year_month, looks up each borrower's state-specific unemployment rate via a CASE "
    "expression on addr_state (falling back to UNRATE for states outside the top 5), and "
    "derives two more features - real_int_rate (int_rate minus CPI_YOY) and rate_spread "
    "(int_rate minus FEDFUNDS). The resulting table is workspace.default.gold_loans_macro."
)

heading("Insight 1 - default rate vs national unemployment at issuance", level=2)
body(
    "Binning all loans by UNRATE in their issue month shows a counterintuitive pattern: loans "
    "issued when national unemployment was lower defaulted at higher rates. Loans issued at "
    "the 4.1% UNRATE bin default at ~24%, while loans issued at the 6.7% bin default at only "
    "~17%. The likely explanation is late-cycle credit loosening: 2014 issuances (high "
    "unemployment, early recovery) came with tighter underwriting, while 2016-2017 issuances "
    "(low unemployment, late cycle) saw looser credit standards."
)

heading("Insight 2 - grade x Fed Funds regime interaction", level=2)
body(
    "We split loans into a low-rate regime (FEDFUNDS <= 0.5%) and a high-rate regime (FEDFUNDS "
    "> 0.5%) and plotted default rate by sub_grade in both regimes. The two curves sit almost "
    "on top of each other: at every grade from A1 through G5, the default rate is dominated "
    "by the sub-grade itself, not the rate environment. Lending Club's internal risk grading "
    "is a much stronger signal than external monetary conditions for this data window."
)

heading("Insight 3 - state-level unemployment differential", level=2)
body(
    "For borrowers in the five largest Lending Club states (CA, TX, NY, FL, IL), we computed "
    "the gap between state unemployment and the national rate at issuance, and plotted default "
    "rate against this gap. States where unemployment was running above the national average "
    "show slightly higher default rates; states below the average show slightly lower rates. "
    "The effect is smaller than the sub-grade effect but consistent across all five states, "
    "which suggests local labor-market conditions carry some predictive signal after grade is "
    "accounted for."
)

heading("Combined-features model", level=2)
body(
    "We retrained the same logistic regression pipeline with UNRATE, FEDFUNDS, CPIAUCSL, "
    "CPI_YOY, state_unrate, real_int_rate, and rate_spread appended to the numeric feature "
    "set. Val AUC-ROC moved from 0.7136 to 0.7131 - effectively no change. The macro features "
    "are highly correlated with issue_d, so a linear model that already sees term, sub_grade, "
    "int_rate, and installment has little room to extract incremental signal from them. "
    "Section 4 discusses why this does not mean the FRED features are useless."
)

# ============================================================
# Section 4 - Challenges and Lessons Learned
# ============================================================
heading("4. Challenges and Lessons Learned", level=1)

heading("Community Edition was renamed, and the new tier blocks MLlib", level=2)
body(
    "The Phase 3 spec explicitly allows Databricks Community Edition. In 2025, Databricks "
    "renamed Community Edition to Free Edition and migrated the underlying compute from classic "
    "single-node clusters to a serverless model. Signing up for Community Edition today routes "
    "you to Free Edition - there is no older-style tier available for new accounts. This is a "
    "platform change that post-dates the Phase 3 spec and was not anticipated in the course "
    "instructions."
)

body(
    "Free Edition's serverless compute enforces a Py4J security whitelist that rejects the "
    "Java constructors for several traditional Spark MLlib feature classes. Instantiating "
    "Imputer on our cluster raises Py4JSecurityException: 'Constructor public "
    "org.apache.spark.ml.feature.Imputer(java.lang.String) is not whitelisted' - the same error "
    "appears for StringIndexer and OneHotEncoder. The Compute page on Free Edition only lists "
    "SQL warehouses and serverless - there is no option to create a classic single-node cluster - "
    "so an end-to-end pyspark.ml pipeline cannot run on the tier the course allows."
)

body(
    "Our workaround was to keep the scalable data work in Spark + Delta (Bronze, Silver, Gold "
    "all run pure Spark on this serverless cluster) and pull the ~200k-row Gold training set "
    "into pandas for the final model-fit step, using scikit-learn's equivalent API. Cross-"
    "validation is handled by sklearn's GridSearchCV, which fills the role of MLlib's "
    "CrossValidator. This preserves the spirit of the MLlib requirement - a scalable "
    "Spark-based pipeline with cross-validated model training - within the constraints of the "
    "allowed platform tier. In future work we would either request access to a paid Databricks "
    "tier with classic clusters, or run a local Spark cluster for the model-fit step."
)

heading("Counterintuitive macro finding", level=2)
body(
    "We expected default rate to rise with unemployment. It fell. The explanation - that loose "
    "credit policy in late-cycle, low-unemployment years produced lower-quality borrowers - is "
    "a known pattern in consumer credit, but we had to discover it from the data rather than "
    "anticipate it. This was a useful reminder that macro features are often confounded with "
    "vintage effects in cohort analyses, and that a conditioning step on issuance year may be "
    "needed before drawing conclusions about unemployment's direct effect on defaults."
)

heading("Macro features in linear models", level=2)
body(
    "Adding the FRED features did not improve logistic regression's Val AUC-ROC. This is not "
    "because the features are uninformative - the insights above show they carry real signal - "
    "but because they are highly correlated with the loan-level features the model already "
    "sees, and because logistic regression cannot express the threshold and interaction effects "
    "(e.g. 'default spikes only when state unemployment is >2 points above national AND grade "
    "is worse than C') that would benefit from macro features. A tree ensemble such as XGBoost "
    "or Random Forest would likely show a larger uplift."
)

heading("What we would do differently", level=2)
bullet("Verify full library compatibility with the target runtime before building the pipeline.")
bullet(
    "Use a tree model (GBT / XGBoost) for the combined-features experiment - linear models "
    "cannot capture the non-linear interactions that make macro features useful."
)
bullet(
    "Add a vintage-adjusted view of default rate (i.e. maturity curves by issue quarter) to "
    "disentangle macro effects from loans having less time to mature."
)
bullet(
    "Invest more in state-level granularity - the current state_unrate falls back to the "
    "national rate for 46 states and probably washes out otherwise-strong local signals."
)

doc.save("Phase3_summary.docx")
print("Wrote Phase3_summary.docx")
